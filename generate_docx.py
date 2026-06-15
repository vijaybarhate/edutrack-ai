from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os

def create_docx():
    doc = Document()
    
    # Title
    title = doc.add_heading('Capstone Project Submission Report: EduTrack AI', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Student Info
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('Sustainable Development Goal: SDG 4 – Quality Education\n')
    run.bold = True
    run = p.add_run('Submitted by: Vijay Barhate\n')
    run.bold = True
    run = p.add_run('Date: June 15, 2026')
    run.bold = True

    # 1. Selected SDG
    doc.add_heading('1. Selected SDG and Reason for Selection', level=1)
    doc.add_paragraph(
        'Goal 4 – Quality Education: Ensure inclusive and equitable quality education and promote '
        'lifelong learning opportunities for all.'
    )
    doc.add_paragraph(
        'Reason for Selection: Education is the foundation of sustainable development. However, many '
        'educational systems rely on "lagging indicators" (final grades) to identify struggling students. '
        'By the time a student fails a final exam, it is often too late for effective intervention. '
        'This issue is particularly severe in overcrowded classrooms or under-resourced schools where '
        'individual performance tracking is manual and slow. I selected SDG 4 because AI can bridge this '
        'gap by providing Early Warning Systems that identify at-risk students during the term, allowing '
        'for timely, data-driven support.'
    )

    # 2. Problem Statement
    doc.add_heading('2. Problem Statement', level=1)
    problem_list = [
        ('What is the problem?', 'Lack of automated, real-time tracking for student academic risk, leading to high failure rates that could have been prevented with early intervention.'),
        ('Where does this problem exist?', 'In schools and universities where large student-to-teacher ratios make it impossible for educators to manually analyze performance trends for every student across multiple subjects.'),
        ('Who is affected?', 'Primarily students at risk of academic failure, but also teachers who lack actionable insights and administrators who cannot effectively allocate support resources.'),
        ('Why is it serious?', 'Academic failure leads to increased dropout rates, loss of motivation, and long-term socio-economic disadvantages, hindering the progress of SDG 4.'),
        ('What happens if it is not solved?', 'Educational inequality will widen as students without private tutoring or individual attention continue to fall through the cracks of a reactive system.')
    ]
    for q, a in problem_list:
        p = doc.add_paragraph(style='List Bullet')
        run = p.add_run(f'{q}: ')
        run.bold = True
        p.add_run(a)

    doc.add_page_break()

    # 3. Proposed Solution
    doc.add_heading('3. Proposed Solution', level=1)
    doc.add_paragraph(
        'EduTrack AI is a full-stack AI platform designed to predict academic risk and automate intervention planning. '
        'It uses a Machine Learning soft-voting ensemble (Random Forest + Logistic Regression) to analyze mid-term performance '
        'data and predict the probability of failure.'
    )
    solution_list = [
        ('How it works', 'The system ingests data like attendance rates, study hours, and mid-term marks. It processes these through a trained classifier that flags students as "High", "Medium", or "Low" risk.'),
        ('Intervention', 'It identifies the "Weakest Subject" for each student and uses rule-based logic to recommend a specific increase in weekly study hours (e.g., +5 hours/week) to mitigate the risk.'),
        ('Impact', 'By converting raw data into "Actionable Insights", it shifts the educational approach from reactive (fixing failures) to proactive (preventing them).'),
        ('Users', 'Teachers (for classroom monitoring), School Administrators (for resource allocation), and Students/Parents (for personalized study plans).')
    ]
    for q, a in solution_list:
        p = doc.add_paragraph(style='List Bullet')
        run = p.add_run(f'{q}: ')
        run.bold = True
        p.add_run(a)

    # 4. Project Features
    doc.add_heading('4. Project Features', level=1)
    features = [
        ('Overview Dashboard', 'Real-time visualization of class-wide metrics (Avg Attendance, Pass Rates) using interactive Plotly charts and KPI cards.'),
        ('Individual Student Lookup', 'A detailed profile search that breaks down a student\'s risk probability per subject and compares their scores to class averages.'),
        ('Risk Prediction Sandbox', 'A simulator where teachers can input hypothetical parameters to see the predicted change in risk.'),
        ('Automated Intervention Plans', 'Logic-driven recommendations that pinpoint the weakest subject and prescribe a personalized study hour adjustment.'),
        ('Report Generation', 'One-click export of student performance records into formatted PDF and CSV files.')
    ]
    for f_title, f_desc in features:
        p = doc.add_paragraph(style='List Bullet')
        run = p.add_run(f'{f_title}: ')
        run.bold = True
        p.add_run(f_desc)

    doc.add_page_break()

    # 5. Technology Stack
    doc.add_heading('5. Technology Stack', level=1)
    techs = ['Python (Core)', 'FastAPI (Backend)', 'Streamlit (Frontend)', 'Scikit-Learn (ML)', 'SQLite (Database)', 'Plotly (Visualization)', 'FPDF2 (PDF Export)']
    doc.add_paragraph(', '.join(techs))

    # 6. Project Screenshots
    doc.add_heading('6. Project Screenshots (Proof of Work)', level=1)
    
    screenshots = [
        ('screenshot_1_revised.png', 'Figure 1: Main Overview Dashboard showing class-wide risk distribution and academic KPIs.'),
        ('screenshot_2_revised.png', 'Figure 2: Individual Student Performance Lookup and risk probability breakdown.'),
        ('screenshot_5_intervention.png', 'Figure 3: Automated Academic Intervention Plan with personalized study hour recommendations.'),
        ('screenshot_6_sandbox_result.png', 'Figure 4: Risk Prediction Sandbox providing real-time probability gauging.'),
        ('screenshot_4_revised.png', 'Figure 5: Dataset Manager for uploading CSV data and triggering automated ML model retraining.')
    ]
    
    for img_path, caption in screenshots:
        if os.path.exists(img_path):
            doc.add_picture(img_path, width=Inches(6.0))
            p = doc.add_paragraph(caption)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            doc.add_paragraph() # Spacer
        else:
            doc.add_paragraph(f'[Image Missing: {img_path}]')

    doc.add_page_break()

    # 7. Open Source Repository
    doc.add_heading('7. Open Source Repository', level=1)
    p = doc.add_paragraph('The complete source code is available on GitHub:\n')
    run = p.add_run('https://github.com/vijaybarhate/edutrack-ai')
    run.font.color.rgb = None # Set color if needed
    run.underline = True

    # 8. Future Scope
    doc.add_heading('8. Future Scope', level=1)
    future_items = [
        'Generative AI Integration: LLM-based tutor (e.g., Google Gemini) for subject-specific help.',
        'Mobile Companion App: Real-time notifications for parents and students.',
        'LMS Connector: Integration with Google Classroom and Moodle.',
        'Advanced Modeling: Incorporate behavioral and socio-economic factors.'
    ]
    for item in future_items:
        doc.add_paragraph(item, style='List Bullet')

    # Save
    filename = 'SDG4_EduTrackAI_VijayBarhate.docx'
    doc.save(filename)
    print(f'Successfully generated {filename}')

if __name__ == '__main__':
    create_docx()
