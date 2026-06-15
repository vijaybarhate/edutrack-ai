import os
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

def create_docx():
    doc = Document()
    
    # Title Page
    title = doc.add_heading('EduTrack AI – SDG 4 Capstone Project', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    p.add_run('Name: ').bold = True
    p.add_run('Vijay Dilip Barhate\n')
    
    p.add_run('Course: ').bold = True
    p.add_run('B.E. Computer Engineering, Saraswati College of Engineering\n')
    
    p.add_run('Date: ').bold = True
    p.add_run('June 15, 2026')
    
    doc.add_page_break()
    
    # Section 1
    doc.add_heading('1. Selected SDG and Reason for Selection', level=1)
    
    p = doc.add_paragraph(style='List Bullet')
    p.add_run('Selected SDG: ').bold = True
    p.add_run('Goal 4 – Quality Education')
    
    p = doc.add_paragraph(style='List Bullet')
    p.add_run('Focus of the SDG: ').bold = True
    p.add_run('SDG 4 focuses on ensuring inclusive and equitable quality education and promoting lifelong learning opportunities for all.')
    
    p = doc.add_paragraph(style='List Bullet')
    p.add_run('Importance of the Issue: ').bold = True
    p.add_run('Educational systems frequently fail to address performance gaps early enough. Student performance declines and attendance gaps go completely unnoticed until it is too late (e.g., after final exams), leading to preventable academic failures and course repetitions. Proactive tracking is critical to maintaining high educational quality and equity.')
    
    # Section 2
    doc.add_heading('2. Problem Statement', level=1)
    
    p = doc.add_paragraph(style='List Bullet')
    p.add_run('What is the problem?: ').bold = True
    p.add_run('Teachers, students, and school administrators lack an automated early warning system for academic risk, preventing timely educational intervention.')
    
    p = doc.add_paragraph(style='List Bullet')
    p.add_run('Where does this problem exist?: ').bold = True
    p.add_run('This issue commonly exists in schools with manual progress-tracking systems and low student-to-teacher ratios, where teachers cannot individually monitor daily performance trends across multiple subjects.')
    
    p = doc.add_paragraph(style='List Bullet')
    p.add_run('Who is affected?: ').bold = True
    p.add_run('The problem directly affects students who are at risk of failing their subjects and dropping out, as well as teachers who lack the tools to perform early interventions.')
    
    p = doc.add_paragraph(style='List Bullet')
    p.add_run('Why is it serious?: ').bold = True
    p.add_run('Without early prediction, performance deficiencies compound. If not solved, this leads to increased course failure rates, higher student dropout percentages, and long-term socio-economic challenges for the affected youth.')

    doc.add_page_break()

    # Section 3
    doc.add_heading('3. Proposed Solution', level=1)
    
    p = doc.add_paragraph(style='List Bullet')
    p.add_run('What the project is: ').bold = True
    p.add_run('EduTrack AI is an intelligent educational assistant that uses machine learning to identify students at risk of academic failure.')
    
    p = doc.add_paragraph(style='List Bullet')
    p.add_run('How the system works: ').bold = True
    p.add_run('The application collects early-term academic data (attendance rate, midterm marks, study hours, past history) and inputs them into a machine learning ensemble (Random Forest + Logistic Regression) that calculates the exact failure risk per subject.')
    
    p = doc.add_paragraph(style='List Bullet')
    p.add_run('How it helps solve the problem: ').bold = True
    p.add_run('It flags students as High, Medium, or Low risk long before final exams, automatically triggering a personalized study plan recommendation to adjust weekly study hours for their weakest subject.')
    
    p = doc.add_paragraph(style='List Bullet')
    p.add_run('Who will use the application: ').bold = True
    p.add_run('School administrators, teachers, and guidance counselors will use the dashboard for class-wide monitoring, while students and parents will receive the individual PDF performance reports.')

    # Section 4
    doc.add_heading('4. Project Features', level=1)
    
    p = doc.add_paragraph(style='List Bullet')
    p.add_run('Overview Dashboard: ').bold = True
    p.add_run('A high-level visual analytics panel displaying KPI cards for class pass rates and average attendance, paired with Plotly charts tracking risk distribution and subject-wise averages to help teachers monitor whole-class performance.')
    
    p = doc.add_paragraph(style='List Bullet')
    p.add_run('Risk Prediction Engine: ').bold = True
    p.add_run('A machine learning-powered sandbox allowing users to input hypothetical attendance and score parameters to instantly calculate failure risk using a Random Forest and Logistic Regression soft-voting ensemble.')
    
    p = doc.add_paragraph(style='List Bullet')
    p.add_run('Personalized Recommendations: ').bold = True
    p.add_run('A rule-based intervention logic that automatically identifies a student\'s weakest subject relative to the class average and calculates the precise addition to weekly study hours needed to mitigate risk.')
    
    p = doc.add_paragraph(style='List Bullet')
    p.add_run('Exportable Reports: ').bold = True
    p.add_run('A reporting tool that allows teachers to instantly generate and download formatted PDF progress reports and CSV raw data containing student grades, risk ratings, and custom action plans.')
    
    doc.add_page_break()

    # Section 5
    doc.add_heading('5. Technology Stack', level=1)
    
    p = doc.add_paragraph(style='List Bullet')
    p.add_run('Programming Language: ').bold = True
    p.add_run('Python (Core application logic and ML modeling)')
    
    p = doc.add_paragraph(style='List Bullet')
    p.add_run('Machine Learning: ').bold = True
    p.add_run('Scikit-Learn (Random Forest Classifier + Logistic Regression voting ensemble, preprocessing pipelines)')
    
    p = doc.add_paragraph(style='List Bullet')
    p.add_run('Backend Framework: ').bold = True
    p.add_run('FastAPI (RESTful API serving prediction, data storage, and model retraining)')
    
    p = doc.add_paragraph(style='List Bullet')
    p.add_run('Frontend Interface: ').bold = True
    p.add_run('Streamlit (Responsive web dashboard and interactive user interface)')
    
    p = doc.add_paragraph(style='List Bullet')
    p.add_run('Database: ').bold = True
    p.add_run('SQLite (Local database for storing student records and tracking academic history)')
    
    p = doc.add_paragraph(style='List Bullet')
    p.add_run('Data Manipulation: ').bold = True
    p.add_run('Pandas and Numpy (Data engineering and processing)')
    
    p = doc.add_paragraph(style='List Bullet')
    p.add_run('Visualization: ').bold = True
    p.add_run('Plotly (Interactive dashboard charts)')
    
    p = doc.add_paragraph(style='List Bullet')
    p.add_run('PDF Generation: ').bold = True
    p.add_run('FPDF2 (Dynamic report layout and export)')

    doc.add_page_break()

    # Section 6
    doc.add_heading('6. Application Screenshots', level=1)
    
    screenshots = [
        ('Overview Dashboard', 'screenshots/overview.png', 'EduTrack AI Overview Dashboard displaying real-time class analytics and at-risk student lists.'),
        ('Individual Lookup', 'screenshots/student_lookup.png', 'Individual Student Lookup interface showing subject breakdown and intervention recommendations.'),
        ('Risk Sandbox', 'screenshots/risk_sandbox.png', 'Risk Prediction Sandbox for simulating student performance scenarios.'),
        ('Dataset Manager', 'screenshots/dataset_manager.png', 'Dataset Manager for uploading school data and retraining the ML model ensemble.')
    ]
    
    for title, img_path, desc in screenshots:
        doc.add_heading(title, level=2)
        
        if os.path.exists(img_path):
            doc.add_picture(img_path, width=Inches(6.0))
        else:
            doc.add_paragraph(f'[Missing Image: {img_path}]')
            
        p = doc.add_paragraph()
        p.add_run(f'{desc}').bold = True
        
        doc.add_paragraph() # Add spacing

    doc.add_page_break()

    # Section 7
    doc.add_heading('7. Open Source Repository', level=1)
    p = doc.add_paragraph()
    p.add_run('GitHub Link: ').bold = True
    p.add_run('https://github.com/vijaybarhate/edutrack-ai')

    # Section 8
    doc.add_heading('8. Future Scope', level=1)
    
    p = doc.add_paragraph(style='List Bullet')
    p.add_run('Chatbot Tutor: ').bold = True
    p.add_run('Integrate a generative AI chatbot tutor using Google Gemini API to provide students with immediate, subject-specific tutoring assistance on their weakest topics.')
    
    p = doc.add_paragraph(style='List Bullet')
    p.add_run('Mobile Application: ').bold = True
    p.add_run('Build a cross-platform mobile companion app for students and parents to view weekly progress reports and study plans on the go.')
    
    p = doc.add_paragraph(style='List Bullet')
    p.add_run('LMS Integration: ').bold = True
    p.add_run('Create integration plugins for standard Learning Management Systems (LMS) like Moodle, Canvas, or Google Classroom to automatically sync attendance and grades.')

    # Section 9
    doc.add_heading('9. Conclusion & Summary', level=1)
    doc.add_paragraph('EduTrack AI demonstrates the power of integrating machine learning and data analytics to serve SDG 4 (Quality Education). By converting static student data into actionable early warning indicators, teachers can intervene before performance gaps widen. The capstone showcases a complete full-stack machine learning pipeline from synthetic dataset generation to API-first serving and interactive visualization, offering an scalable, automated approach to supporting student success.')

    # Save
    filename = 'SDG4_EduTrackAI_VijayBarhate.docx'
    doc.save(filename)
    print(f'Successfully generated {filename}')

if __name__ == '__main__':
    create_docx()
